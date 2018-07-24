# -*- coding: utf-8 -*
# author luxin

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENTATION
import winreg
import os
import re
from service.utils import conf


# 创建word文档工具类
class WordUtils(object):
    def __init__(self):
        self.document = Document(conf.word_path())
        self.paragraph = None
        self.table = None

    # 创建段落并使用样式
    def add_paragraph(self, words, style_name):
        style = self.__check_style(style_name)
        self.paragraph = self.document.add_paragraph(words, style=style)

    # 向段落中追加数据
    def add_to_paragraph(self, words, **font_sets):
        paragraph = self.paragraph
        if paragraph is not None:
            run = paragraph.add_run(words)
            dict = {'name': self.__set_font_name, 'size': self.__set_font_size, 'bold': self.__set_font_bold,
                    'italic': self.__set_font_italic,
                    'underline': self.__set_font_underline}
            for key in font_sets.keys():
                key = key.lower()
                dict[key](run, font_sets[key])

    # 添加图片
    def add_picture(self, src, width):
        self.document.add_picture(src, Pt(width))

    # 创建表格
    def add_table(self, rows, cols, style_name):
        style = self.__check_table_style(style_name)
        if style is not None:
            self.table = self.document.add_table(rows, cols, style)
        else:
            self.table = self.document.add_table(rows, cols)

    # 合并单元格
    def merge_cell(self, start_row, start_col, end_row, end_col):
        self.table.cell(start_row, start_col).merge(self.table.cell(end_row, end_col))

    # 填充单元格
    def insert_cell(self, row, col, words):
        paragraphs = self.table.cell(row, col).paragraphs
        last_index = paragraphs.__len__() - 1
        self.paragraph = paragraphs[last_index]
        self.paragraph.add_run(words)

    # 单元格新增段落
    def cell_add_paragraph(self, row, col):
        self.table.cell(row, col).add_paragraph()

    # 检查样式，存在返回当前样式，否则返回Normal样式
    def __check_style(self, style_name):
        if style_name in self.document.styles:
            return self.document.styles[style_name]
        else:
            return self.document.styles[u'Normal']
        # 检查样式，存在返回当前样式，否则返回Normal样式

    def __check_table_style(self, style_name):
        if style_name in self.document.styles:
            style = self.document.styles[style_name]
            if style.type == WD_STYLE_TYPE.TABLE:
                return style
        for style in self.document.styles:
            if style.type == WD_STYLE_TYPE.TABLE:
                return style
        return None

    # 纸张横向
    def section_landscape(self):
        length = self.document.sections.__len__() - 1
        section = self.document.sections[length - 1]
        new_width, new_height = section.page_height, section.page_width
        section = self.document.add_section()
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

    # 纸张纵向
    def section_portrait(self):
        length = self.document.sections.__len__() - 1
        section = self.document.sections[length - 1]
        new_width, new_height = section.page_height, section.page_width
        section = self.document.add_section()
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width = new_height
        section.page_height = new_width

    # 保存word文档
    def save_document(self):
        if self.document is not None:
            # self.document.save(self.__get_file_full_name())
            from service.utils import conf
            self.document.save(conf.doc_path)
        self.document = None
        self.paragraph = None
        self.table = None

    # 获取文件全路径
    def __get_file_full_name(self):
        words = "\demo"
        file_path = self.__get_desktop()
        full_name = file_path + words + '.doc'
        count = 0
        while os.path.exists(full_name):
            count = count + 1
            full_name = file_path + words + '(' + str(count) + ').doc'
        return full_name

    # 获取桌面路径
    def __get_desktop(self):
        key = winreg.OpenKey(winreg.HKEY_CURRENT_USER,
                             r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders')
        return winreg.QueryValueEx(key, "Desktop")[0]

    # 设置字体名称
    def __set_font_name(self, run, val):
        run.font.name = val
        if self.__contain_zh(val):
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), val)

    # 设置字体字号（单位：磅）
    def __set_font_size(self, run, val):
        run.font.size = Pt(val)

    # 设置字体加粗
    def __set_font_bold(self, run, val):
        run.bold = val

    # 设置字体斜体
    def __set_font_italic(self, run, val):
        run.italic = val

    # 设置字体下划线
    def __set_font_underline(self, run, val):
        run.underline = val

    # 判断是否含有中文
    def __contain_zh(self, words):
        zh_pattern = re.compile(u'[\u4e00-\u9fa5]+')
        match = zh_pattern.search(words)
        if match:
            return True
        else:
            return False
